# app.py - UNIFIED INTERFACE VERSION WITH ALL CRITICAL ISSUES RESOLVED

# STEP 1: Set debug mode control FIRST
import os
import sys
import warnings
import tempfile
import hashlib
from pathlib import Path
import streamlit as st
import time
import datetime
from PIL import Image
import io

# Optional imports for document processing
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import psutil
    SYSTEM_MONITOR = True
except ImportError:
    SYSTEM_MONITOR = False

# Set default debug mode to FALSE unless explicitly enabled
if 'DEBUG_MODE' not in os.environ:
    os.environ['DEBUG_MODE'] = 'false'

# Global flags to prevent repeated debug messages
_SQLITE_LOGGED = False
_ENV_SETUP_LOGGED = False

# ===== CRITICAL FIX #1: GLOBAL UI STATE INITIALIZATION =====
def initialize_global_ui_state():
    """Initialize global UI state - SAFE TO CALL MULTIPLE TIMES"""
    if 'ui_state' not in st.session_state:
        st.session_state.ui_state = {
            'current_result': None,
            'current_question': "",
            'last_asked_question': "",
            'show_sources': False,
            'show_analytics': False,
            'question_counter': 0,
            'initialized': True
        }
    
    # Initialize other session state variables if needed
    if 'show_doc_findings' not in st.session_state:
        st.session_state.show_doc_findings = False
    if 'show_doc_sources' not in st.session_state:
        st.session_state.show_doc_sources = False
    if 'show_img_details' not in st.session_state:
        st.session_state.show_img_details = False

# Call global initialization immediately
initialize_global_ui_state()

def debug_log(message, once_only=False, key=None):
    """Only print debug messages if debug mode is enabled, with optional one-time logging"""
    if os.getenv('DEBUG_MODE', 'false').lower() == 'true':
        if once_only and key:
            if 'debug_logged' not in st.session_state:
                st.session_state.debug_logged = set()
            
            if key not in st.session_state.debug_logged:
                print(f"DEBUG: {message}")
                st.session_state.debug_logged.add(key)
        elif not once_only:
            print(f"DEBUG: {message}")

def debug_timing(func_name, start_time):
    """Debug timing helper"""
    if os.getenv('DEBUG_MODE', 'false').lower() == 'true':
        elapsed = time.time() - start_time
        print(f"TIMING: {func_name} took {elapsed:.3f}s")

# STEP 2: Suppress ALL Google/gRPC warnings
os.environ['GRPC_VERBOSITY'] = 'ERROR'
os.environ['GLOG_minloglevel'] = '3'
os.environ['TF_CPP_MIN_LOG_LEVEL'] = '3'
os.environ['GRPC_ENABLE_FORK_SUPPORT'] = '0'
os.environ['GRPC_POLL_STRATEGY'] = 'poll'

# Suppress Python warnings
warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", message=".*absl.*")
warnings.filterwarnings("ignore", message=".*gRPC.*")
warnings.filterwarnings("ignore", module="langchain")

# STEP 3: SQLite3 Fix with SILENT logging
def setup_sqlite_with_clean_logging():
    """Setup SQLite with one-time logging only"""
    global _SQLITE_LOGGED
    
    try:
        __import__('pysqlite3')
        sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
        if not _SQLITE_LOGGED:
            debug_log("Using pysqlite3", once_only=True, key="sqlite_pysqlite")
            _SQLITE_LOGGED = True
    except ImportError:
        if not _SQLITE_LOGGED:
            debug_log("Using system sqlite3", once_only=True, key="sqlite_system")
            _SQLITE_LOGGED = True

# STEP 4: Regular imports
from rag_system import HybridRAGSystem as EnhancedRAGSystem, create_rag_system

# Safe import for compliance features
try:
    from compliance_analyzer import ComplianceAnalyzer, create_compliance_interface
    COMPLIANCE_FEATURES_AVAILABLE = True
except ImportError:
    COMPLIANCE_FEATURES_AVAILABLE = False
    debug_log("Compliance features not available")

# Page configuration
st.set_page_config(
    page_title="Children's Home Management System",
    page_icon="🏠",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ===== CACHED CSS =====
@st.cache_data
def get_app_css():
    return """
    <style>
    .main-header {
        text-align: center;
        padding: 1rem 0 2rem 0;
        border-bottom: 1px solid #f0f0f0;
        margin-bottom: 2rem;
    }
    .testing-banner {
        background: linear-gradient(90deg, #ff6b6b, #feca57);
        color: white;
        padding: 0.5rem;
        text-align: center;
        font-weight: bold;
        margin-bottom: 1rem;
        border-radius: 5px;
    }
    .result-container {
        background-color: #f8f9fa;
        border: 1px solid #dee2e6;
        border-radius: 8px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    .question-container {
        background-color: #e3f2fd;
        border-left: 4px solid #2196f3;
        padding: 1rem;
        margin: 1rem 0;
        border-radius: 0 8px 8px 0;
    }
    .clean-separator {
        margin: 2rem 0;
        border-top: 2px solid #e9ecef;
    }
    .stTextArea > div > div > textarea {
        border-radius: 8px;
    }
    .stButton > button {
        border-radius: 8px;
    }
    </style>
    """

# ===== AUTHENTICATION SYSTEM =====
def get_tester_credentials():
    """Get tester credentials with fallback system"""
    try:
        testers = st.secrets.get("testers", {})
        if testers:
            return testers
    except Exception as e:
        debug_log(f"Secrets access failed: {e}")
    
    # Fallback for local development
    return {
        "DEMO001": {
            "password": "DemoAccess2024!",
            "name": "Demo Tester",
            "email": "demo@example.com",
            "expires": "2025-12-31"
        },
        "TEST001": {
            "password": "TestAccess456!",
            "name": "Beta Tester 1", 
            "email": "tester1@example.com",
            "expires": "2025-12-31"
        }
    }

def check_session_valid():
    """Check if current session is still valid"""
    if not st.session_state.get('authenticated_tester', False):
        return False
    
    session_start = st.session_state.get('session_start', 0)
    current_time = time.time()
    session_duration = current_time - session_start
    
    if session_duration > 7200:  # 2 hours
        st.session_state['authenticated_tester'] = False
        return False
    
    return True

def log_tester_activity(tester_id, action, details=""):
    """Log tester activity with completely clean output"""
    try:
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = {
            "timestamp": timestamp,
            "tester_id": tester_id or "anonymous",
            "action": action,
            "details": details
        }
        
        if 'activity_log' not in st.session_state:
            st.session_state['activity_log'] = []
        
        st.session_state['activity_log'].append(log_entry)
        
        critical_actions = [
            'successful_login', 'failed_login', 'logout',
            'system_error', 'feedback_submitted'
        ]
        
        if action in critical_actions:
            print(f"TESTER: {timestamp} - {tester_id or 'anonymous'} - {action}")
            if action == 'system_error' and details:
                print(f"  └─ {details}")
        elif os.getenv('DEBUG_MODE', 'false').lower() == 'true':
            print(f"DEBUG: {timestamp} - {tester_id or 'anonymous'} - {action}")
            if details:
                print(f"  └─ {details}")
        
    except Exception:
        pass

def validate_and_authenticate(tester_id, access_code):
    """Validate credentials and set up authenticated session"""
    try:
        if not tester_id or not access_code:
            return False
        
        valid_testers = get_tester_credentials()
        
        if tester_id not in valid_testers:
            return False
        
        tester_info = valid_testers[tester_id]
        if tester_info.get("password") != access_code:
            return False
        
        expiry_date = tester_info.get("expires", "2099-12-31")
        try:
            expiry = datetime.datetime.strptime(expiry_date, "%Y-%m-%d").date()
            if datetime.date.today() > expiry:
                return False
        except:
            pass
        
        st.session_state['authenticated_tester'] = True
        st.session_state['tester_id'] = tester_id
        st.session_state['tester_info'] = tester_info
        st.session_state['session_start'] = time.time()
        
        log_tester_activity(tester_id, "successful_login", f"User: {tester_info.get('name', 'Unknown')}")
        
        return True
        
    except Exception as e:
        print(f"ERROR: Authentication failed: {e}")
        return False

def show_login_interface():
    """Streamlined login interface"""
    st.markdown(get_app_css(), unsafe_allow_html=True)
    
    st.markdown("""
    <div class="login-container" style="max-width: 600px; margin: 2rem auto; padding: 2rem; background: #f8f9fa; border-radius: 10px; border: 1px solid #e9ecef;">
    <div style="text-align: center; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); color: white; padding: 2rem; border-radius: 10px; margin-bottom: 2rem;">
        <h1>🏠 Children's Home Management System</h1>
        <h3>Beta Testing Portal</h3>
        <p>Secure access for authorized testers</p>
    </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("### 🔐 Tester Login")
    
    tester_id = st.text_input(
        "Tester ID:",
        placeholder="e.g., DEMO001",
        help="Your unique tester identifier",
        key="login_tester_id_input"
    )
    
    access_code = st.text_input(
        "Access Code:",
        type="password",
        placeholder="Enter your access code",
        help="Secure password provided with your Tester ID",
        key="login_access_code_input"
    )
    
    if st.button("🚀 Start Testing Session", type="primary", use_container_width=True, key="login_submit_button"):
        if validate_and_authenticate(tester_id, access_code):
            st.success("✅ Authentication successful!")
            time.sleep(1)
            st.rerun()
        else:
            st.error("❌ Invalid credentials. Please verify your Tester ID and Access Code.")
            log_tester_activity(tester_id or "unknown", "failed_login", "Invalid credentials")
    
    with st.expander("ℹ️ Testing Information"):
        st.markdown("""
        **Session Details:**
        - Duration: 2 hours per session
        - Auto-logout after inactivity
        - Activity logging for improvement
        
        **Need Access?**
        Contact the administrator to receive your testing credentials.
        """)

def show_session_info():
    """Show session information in sidebar"""
    with st.sidebar:
        st.markdown("---")
        st.markdown("### 👤 Session Info")
        
        tester_info = st.session_state.get('tester_info', {})
        st.write(f"**Tester:** {tester_info.get('name', 'Unknown')}")
        st.write(f"**ID:** {st.session_state.get('tester_id', 'N/A')}")
        
        session_start = st.session_state.get('session_start', time.time())
        elapsed = time.time() - session_start
        remaining = max(0, 7200 - elapsed)
        
        hours = int(remaining // 3600)
        minutes = int((remaining % 3600) // 60)
        st.write(f"**Time remaining:** {hours}h {minutes}m")
        
        if st.button("🚪 End Session", type="secondary", key="sidebar_logout_button"):
            log_tester_activity(st.session_state.get('tester_id', 'unknown'), "logout")
            
            for key in ['authenticated_tester', 'tester_id', 'tester_info', 'session_start']:
                if key in st.session_state:
                    del st.session_state[key]
            
            st.rerun()

def require_authentication():
    """Main authentication function"""
    if check_session_valid():
        show_session_info()
        return True
    
    st.session_state['authenticated_tester'] = False
    show_login_interface()
    return False

# ===== ENVIRONMENT SETUP =====
def setup_environment_variables():
    """Enhanced environment setup with better local support"""
    global _ENV_SETUP_LOGGED
    
    try:
        try:
            openai_key = st.secrets.get("OPENAI_API_KEY")
            google_key = st.secrets.get("GOOGLE_API_KEY")
            
            if openai_key and google_key:
                os.environ['OPENAI_API_KEY'] = openai_key
                os.environ['GOOGLE_API_KEY'] = google_key
                if not _ENV_SETUP_LOGGED:
                    debug_log("API keys loaded from Streamlit secrets", once_only=True, key="env_streamlit")
                    _ENV_SETUP_LOGGED = True
                return True
        except Exception as e:
            debug_log(f"Streamlit secrets failed: {e}")
        
        try:
            api_keys = st.secrets.get("api_keys", {})
            openai_key = api_keys.get("openai")
            google_key = api_keys.get("google")
            
            if openai_key and google_key:
                os.environ['OPENAI_API_KEY'] = openai_key
                os.environ['GOOGLE_API_KEY'] = google_key
                if not _ENV_SETUP_LOGGED:
                    debug_log("API keys loaded from api_keys section", once_only=True, key="env_nested")
                    _ENV_SETUP_LOGGED = True
                return True
        except Exception as e:
            debug_log(f"Nested secrets failed: {e}")
        
        try:
            import toml
            secrets_path = Path(".streamlit/secrets.toml")
            
            if secrets_path.exists():
                with open(secrets_path, 'r') as f:
                    secrets = toml.load(f)
                
                openai_key = secrets.get("OPENAI_API_KEY")
                google_key = secrets.get("GOOGLE_API_KEY")
                
                if not (openai_key and google_key):
                    api_keys = secrets.get("api_keys", {})
                    openai_key = api_keys.get("openai")
                    google_key = api_keys.get("google")
                
                if openai_key and google_key:
                    os.environ['OPENAI_API_KEY'] = openai_key
                    os.environ['GOOGLE_API_KEY'] = google_key
                    if not _ENV_SETUP_LOGGED:
                        debug_log("API keys loaded via manual TOML parsing", once_only=True, key="env_manual")
                        _ENV_SETUP_LOGGED = True
                    return True
        except Exception as e:
            debug_log(f"Manual TOML loading failed: {e}")
        
        try:
            openai_key = os.getenv("OPENAI_API_KEY")
            google_key = os.getenv("GOOGLE_API_KEY")
            
            if openai_key and google_key:
                if not _ENV_SETUP_LOGGED:
                    debug_log("API keys found in environment variables", once_only=True, key="env_vars")
                    _ENV_SETUP_LOGGED = True
                return True
        except Exception as e:
            debug_log(f"Environment variables check failed: {e}")
        
        debug_log("WARNING: No API keys found in any location")
        return False
        
    except Exception as e:
        debug_log(f"Environment setup failed: {e}")
        return False

# ===== RAG SYSTEM INITIALIZATION =====
@st.cache_resource
def initialize_rag_system():
    """Initialize RAG system with enhanced error handling"""
    debug_log("Initializing RAG system...", once_only=True, key="rag_init_start")
    
    try:
        if not os.path.exists("faiss_index"):
            st.error("❌ FAISS index not found!")
            st.info("💡 The vector database needs to be created first.")
            st.info("🔧 Please run the ingestion script to create the index.")
            return None
        
        import asyncio
        
        os.environ['GRPC_ENABLE_FORK_SUPPORT'] = '0'
        os.environ['GRPC_POLL_STRATEGY'] = 'poll'
        
        try:
            loop = asyncio.get_event_loop()
            if loop.is_closed():
                asyncio.set_event_loop(asyncio.new_event_loop())
        except RuntimeError:
            asyncio.set_event_loop(asyncio.new_event_loop())
        
        if sys.platform.startswith('win'):
            if hasattr(asyncio, 'WindowsProactorEventLoopPolicy'):
                asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())
        
        debug_log("Asyncio configured", once_only=True, key="asyncio_config")
        
    except Exception as e:
        debug_log(f"Asyncio configuration failed: {e}")
    
    try:
        rag_system = create_rag_system()
        debug_log("RAG system initialized successfully", once_only=True, key="rag_init_success")
        return rag_system
        
    except ImportError as e:
        st.error(f"❌ Import Error: {e}")
        st.info("💡 Missing required packages. Check requirements.txt")
        return None
        
    except Exception as e:
        st.error(f"❌ RAG System Error: {e}")
        debug_log(f"RAG system initialization failed: {e}")
        
        error_msg = str(e).lower()
        if "faiss" in error_msg:
            st.info("💡 FAISS vector database not found or corrupted")
        elif "api" in error_msg or "key" in error_msg:
            st.info("💡 Check API key configuration in secrets")
        elif "import" in error_msg or "module" in error_msg:
            st.info("💡 Check that all dependencies are in requirements.txt")
        
        return None

# ===== PERFORMANCE MODE SELECTOR =====
def add_performance_mode_selector():
    """Add performance mode selector to sidebar"""
    with st.sidebar:
        st.markdown("---")
        st.markdown("### ⚡ Performance Settings")
        
        mode_descriptions = {
            "fast": "🚀 Fast (2-4 docs, ~5s response)",
            "balanced": "⚖️ Balanced (3-7 docs, ~8s response)", 
            "comprehensive": "🔍 Comprehensive (5-12 docs, ~15s response)"
        }
        
        current_mode = st.session_state.get('performance_mode', 'balanced')
        
        selected_mode = st.selectbox(
            "Response Mode:",
            ["fast", "balanced", "comprehensive"],
            index=["fast", "balanced", "comprehensive"].index(current_mode),
            format_func=lambda x: mode_descriptions[x],
            help="Choose speed vs comprehensiveness trade-off",
            key="performance_mode_selector_widget"
        )
        
        st.session_state['performance_mode'] = selected_mode
        
        if selected_mode == "fast":
            st.caption("⚡ Optimized for speed. Best for simple, factual questions.")
        elif selected_mode == "balanced":
            st.caption("⚖️ Good balance of speed and completeness. Recommended for most queries.")
        else:
            st.caption("🔍 Maximum thoroughness. Best for complex strategic questions.")
        
        return selected_mode

# ===== UI STATE MANAGEMENT =====
def ensure_ui_state():
    """Defensive UI state initialization - SAFE TO CALL ANYWHERE"""
    initialize_global_ui_state()

def clear_ui_state():
    """Clear UI state for new question"""
    ensure_ui_state()
    st.session_state.ui_state.update({
        'current_result': None,
        'current_question': "",
        'last_asked_question': "",
        'show_sources': False,
        'show_analytics': False,
        'question_counter': st.session_state.ui_state.get('question_counter', 0) + 1
    })

# ===== SOURCE PROCESSING FUNCTIONS =====
def extract_source_info(source):
    """Extract source information from Document objects or dictionaries"""
    # Handle Document objects from LangChain
    if hasattr(source, 'metadata') and hasattr(source, 'page_content'):
        metadata = source.metadata
        return {
            'title': metadata.get('title', ''),
            'source': metadata.get('source', ''),
            'source_type': metadata.get('source_type', ''),
            'word_count': len(source.page_content.split()) if source.page_content else 0,
            'page_content': source.page_content
        }
    # Handle dictionary sources
    elif isinstance(source, dict):
        return {
            'title': source.get('title', ''),
            'source': source.get('source', ''),
            'source_type': source.get('source_type', ''),
            'word_count': source.get('word_count', 0),
            'page_content': source.get('page_content', '')
        }
    else:
        return {
            'title': f'Unknown Source',
            'source': '',
            'source_type': 'unknown',
            'word_count': 0,
            'page_content': ''
        }

def deduplicate_sources(sources):
    """Deduplicate sources by filename/title"""
    unique_sources = {}
    
    for source in sources:
        source_info = extract_source_info(source)
        
        # Create a unique key for this document
        source_file = source_info['source']
        title = source_info['title']
        
        if source_file:
            filename = os.path.basename(source_file)
            doc_key = filename
        elif title and len(title.strip()) > 3:
            doc_key = title.strip()
        else:
            content = source_info['page_content']
            doc_key = f"doc_{hash(content[:100]) % 1000}"
        
        # Only keep first occurrence of each document
        if doc_key not in unique_sources:
            unique_sources[doc_key] = {
                'title': title,
                'source_file': source_file,
                'source_type': source_info['source_type'],
                'word_count': source_info['word_count'],
                'chunks_found': 1
            }
        else:
            unique_sources[doc_key]['chunks_found'] += 1
            # Update word count if it's higher
            if source_info['word_count'] > unique_sources[doc_key]['word_count']:
                unique_sources[doc_key]['word_count'] = source_info['word_count']
    
    return unique_sources

def clean_source_title(title, source_file, index):
    """Clean up source title for display"""
    if not title or len(title.strip()) < 3:
        if source_file:
            filename = os.path.basename(source_file)
            clean_title = filename.replace('.pdf', '').replace('.docx', '').replace('.txt', '')
            clean_title = clean_title.replace('_', ' ').replace('-', ' ')
            title = ' '.join(word.capitalize() for word in clean_title.split())
        else:
            title = f'Reference Document {index}'
    
    # Truncate very long titles
    if len(title) > 80:
        title = title[:77] + "..."
    
    return title

# ===== UNIFIED INTERFACE FUNCTIONS =====
def show_unified_input_interface():
    """Single unified input interface for all interactions"""
    ensure_ui_state()
    
    st.subheader("💬 Ask Your Question")
    
    # Question input FIRST (above file upload)
    question_key = f"unified_question_input_{st.session_state.ui_state.get('question_counter', 0)}"
    
    user_question = st.text_area(
        "Describe your situation or ask any question:",
        value="",
        placeholder="Examples:\n• 'What should we focus on for our next inspection?'\n• 'How do we handle a safeguarding concern?'\n• 'Review our medication policy for compliance'\n• 'Analyze this kitchen photo for safety issues'",
        height=120,
        key=question_key,
        help="Ask about operations, compliance, policies, or any children's home management topic"
    )
    
    # File upload area SECOND (below question, compact design)
    st.markdown("---")
    
    # Create two columns to make file upload more compact
    col1, col2 = st.columns([2, 1])
    
    with col1:
        uploaded_files = st.file_uploader(
            "📎 Optional: Upload files for analysis",
            type=['pdf', 'docx', 'txt', 'md', 'png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            help="Upload documents or images to analyze with your question",
            key="unified_file_uploader"
        )
    
    with col2:
        if uploaded_files:
            st.markdown("**Files Ready:**")
            for file in uploaded_files:
                file_type = "📄" if file.type.startswith(('application', 'text')) else "📷"
                st.write(f"{file_type} {file.name}")
        else:
            st.markdown("**Supported:**")
            st.write("📄 PDF, Word, Text")
            st.write("📷 Images (PNG, JPG)")
    
    # Submit button
    st.markdown("---")
    if st.button("🧠 Get Expert Guidance", type="primary", use_container_width=True, key="unified_submit_button"):
        if user_question.strip() or uploaded_files:
            process_unified_request(user_question, uploaded_files)
        else:
            st.warning("⚠️ Please enter a question or upload files for analysis")
    
    # Quick actions (simplified)
    show_quick_actions()

def show_quick_actions():
    """Quick action buttons with unique keys"""
    st.markdown("---")
    st.markdown("### 🚀 Quick Start Options")
    
    quick_actions = [
        ("🚨 Inspection Prep", "We have an upcoming Ofsted inspection. What should we focus on to ensure we're fully prepared?"),
        ("👥 Staff Issues", "We're experiencing staff retention challenges. What strategies can we implement?"),
        ("🏠 New Admission", "We're admitting a new child. What processes should we prioritize for a successful transition?"),
        ("📋 Policy Review", "We need to review and update our policies. What current best practices should we include?"),
        ("💰 Budget Planning", "Help us develop an effective budget strategy for the upcoming financial year."),
        ("🎯 Quality Improvement", "We want to enhance our care quality and move towards outstanding. What should we focus on?")
    ]
    
    cols = st.columns(3)
    for i, (title, question) in enumerate(quick_actions):
        with cols[i % 3]:
            # UNIQUE KEY for each quick action button
            if st.button(title, key=f"quick_action_button_{i}", use_container_width=True):
                process_unified_request(question)

def process_unified_request(user_question, uploaded_files=None):
    """Process unified request with BOTH document and image analysis support"""
    ensure_ui_state()
    timing_start = time.time()
    
    performance_mode = st.session_state.get('performance_mode', 'balanced')
    
    # Separate files by type
    image_files = []
    document_files = []
    
    if uploaded_files:
        for file in uploaded_files:
            file_extension = file.name.split('.')[-1].lower()
            if file_extension in ['png', 'jpg', 'jpeg']:
                image_files.append(file)
            elif file_extension in ['pdf', 'docx', 'txt', 'md']:
                document_files.append(file)
    
    # Process documents for context (existing logic)
    document_context = ""
    if document_files:
        document_context = process_uploaded_documents(document_files)
    
    # Build the question with document context
    if document_context:
        full_question = f"{user_question}\n\nDocument Content:\n{document_context}"
        is_file_analysis = True
    else:
        full_question = user_question
        is_file_analysis = bool(image_files)  # True if images uploaded
    
    # Create default question if none provided but files uploaded
    if not user_question.strip() and (image_files or document_files):
        if image_files and document_files:
            user_question = "Analyze these documents and images for compliance, safety, and facility management."
            full_question = f"{user_question}\n\nDocument Content:\n{document_context}"
        elif image_files:
            user_question = "Analyze these images for safety, compliance, and facility management."
            full_question = user_question
        elif document_files:
            user_question = "Analyze these documents for key findings, compliance, and recommendations."
            full_question = f"{user_question}\n\nDocument Content:\n{document_context}"
    
    log_tester_activity(
        st.session_state.get('tester_id', 'unknown'),
        "unified_request_submitted",
        f"Mode: {performance_mode}, Docs: {len(document_files)}, Images: {len(image_files)}, Question: {user_question[:50]}..."
    )
    
    progress_placeholder = st.empty()
    
    try:
        with progress_placeholder:
            if image_files and document_files:
                st.info(f"🔍 Analyzing {len(document_files)} document(s) and {len(image_files)} image(s)...")
            elif image_files:
                st.info(f"📷 Analyzing {len(image_files)} image(s) with AI vision...")
            elif document_files:
                st.info(f"📄 Analyzing {len(document_files)} document(s)...")
            else:
                st.info("🔍 Processing your question...")
        
        # Call enhanced RAG system with both document context and images
        result = st.session_state.rag_system.query(
            question=full_question,
            k=5,
            response_style="standard",
            performance_mode="balanced",
            is_file_analysis=is_file_analysis,
            uploaded_files=uploaded_files, 
            uploaded_images=uploaded_images if 'uploaded_images' in locals() else None # Vision AI handles images
            # Document context is already in full_question, so RAG can find related docs
        )
        
        progress_placeholder.empty()
        
        if result and result.get("answer"):
            st.session_state.ui_state.update({
                'current_result': result,
                'last_asked_question': user_question,
                'show_sources': False,
                'show_analytics': False
            })
            
            # Enhanced success message
            success_parts = []
            if document_files:
                success_parts.append(f"{len(document_files)} document(s)")
            if image_files:
                vision_model = result.get("metadata", {}).get("vision_model", "AI vision")
                success_parts.append(f"{len(image_files)} image(s) with {vision_model}")
            
            if success_parts:
                st.success(f"✅ Analysis complete! Processed {' and '.join(success_parts)}")
            else:
                st.success("✅ Analysis complete!")

            st.rerun()
            
            # Display results
            show_clean_result_display()
            
            log_tester_activity(
                st.session_state.get('tester_id', 'unknown'),
                "successful_response",
                f"Mode: {performance_mode}, Docs: {len(document_files)}, Images: {len(image_files)}, Vision: {result.get('metadata', {}).get('vision_analysis_performed', False)}, Time: {time.time() - timing_start:.1f}s"
            )
            
        else:
            st.error("❌ Sorry, I couldn't generate a response.")
            
    except Exception as e:
        progress_placeholder.empty()
        st.error("❌ An error occurred while processing your request.")
        print(f"ERROR: {str(e)}")
        
        log_tester_activity(
            st.session_state.get('tester_id', 'unknown'),
            "system_error",
            f"Mode: {performance_mode}, Docs: {len(document_files)}, Images: {len(image_files)}, Error: {str(e)[:100]}"
        )


def process_uploaded_documents(uploaded_files):
    """Process uploaded DOCUMENTS (PDF, DOCX, TXT, MD) for text analysis"""
    file_contents = []
    
    for file in uploaded_files:
        try:
            file_extension = file.name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                if PDF_SUPPORT:
                    import PyPDF2
                    file.seek(0)  # Reset file pointer
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
                    file_contents.append(f"DOCUMENT: {file.name}\n{content}")
                else:
                    file_contents.append(f"DOCUMENT: {file.name} (PDF processing not available)")
            
            elif file_extension == 'docx':
                if DOCX_SUPPORT:
                    import docx
                    file.seek(0)  # Reset file pointer
                    doc = docx.Document(file)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    file_contents.append(f"DOCUMENT: {file.name}\n{content}")
                else:
                    file_contents.append(f"DOCUMENT: {file.name} (Word processing not available)")
            
            elif file_extension in ['txt', 'md']:
                file.seek(0)  # Reset file pointer
                content = file.read().decode('utf-8')
                file_contents.append(f"DOCUMENT: {file.name}\n{content}")
                
        except Exception as e:
            file_contents.append(f"ERROR processing {file.name}: {str(e)}")
    
    return "\n\n---\n\n".join(file_contents)

def process_uploaded_files(uploaded_files):
    """Process uploaded DOCUMENTS only - images handled by vision AI"""
    file_contents = []
    
    for file in uploaded_files:
        try:
            file_extension = file.name.split('.')[-1].lower()
            
            # Skip images - they're handled by vision AI now
            if file_extension in ['png', 'jpg', 'jpeg']:
                continue
                
            elif file_extension == 'pdf':
                if PDF_SUPPORT:
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in pdf_reader.pages:
                        content += page.extract_text() + "\n"
                    file_contents.append(f"DOCUMENT: {file.name}\n{content}")
                else:
                    file_contents.append(f"DOCUMENT: {file.name} (PDF processing not available)")
            
            elif file_extension == 'docx':
                if DOCX_SUPPORT:
                    import docx
                    doc = docx.Document(file)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    file_contents.append(f"DOCUMENT: {file.name}\n{content}")
                else:
                    file_contents.append(f"DOCUMENT: {file.name} (Word processing not available)")
            
            elif file_extension in ['txt', 'md']:
                content = file.read().decode('utf-8')
                file_contents.append(f"DOCUMENT: {file.name}\n{content}")
                
        except Exception as e:
            file_contents.append(f"ERROR processing {file.name}: {str(e)}")
    
    return "\n\n---\n\n".join(file_contents)

def determine_image_context(filename):
    """Determine the likely context of an uploaded image for better analysis"""
    filename_lower = filename.lower()
    
    # Kitchen/dining related
    if any(keyword in filename_lower for keyword in ['kitchen', 'dining', 'food', 'cook']):
        return "Kitchen/dining area - analyzing for food safety and hygiene compliance"
    
    # Bedroom/living areas
    elif any(keyword in filename_lower for keyword in ['bedroom', 'living', 'lounge', 'room']):
        return "Living space - analyzing for safety, homeliness, and child-friendly environment"
    
    # Bathroom areas
    elif any(keyword in filename_lower for keyword in ['bathroom', 'toilet', 'shower']):
        return "Bathroom facility - analyzing for privacy, safety, and hygiene standards"
    
    # Outdoor areas
    elif any(keyword in filename_lower for keyword in ['garden', 'outdoor', 'playground', 'yard']):
        return "Outdoor space - analyzing for safety, security, and recreational value"
    
    # Office/admin areas
    elif any(keyword in filename_lower for keyword in ['office', 'admin', 'staff']):
        return "Administrative area - analyzing for professional standards and security"
    
    # General facility
    elif any(keyword in filename_lower for keyword in ['facility', 'home', 'building', 'entrance']):
        return "General facility - analyzing for overall safety, compliance, and environment quality"
    
    else:
        return "Facility image - analyzing for safety, compliance, and quality standards"

# ===== RESULT DISPLAY FUNCTIONS =====
def show_clean_result_display():
    """Clean result display with question box and no input interface"""
    ensure_ui_state()
    result = st.session_state.ui_state['current_result']
    question = st.session_state.ui_state['last_asked_question']
    
    # Question display in clean container (as before)
    st.markdown("""
    <div class="question-container">
        <h4>🤔 Your Question:</h4>
        <p>{}</p>
    </div>
    """.format(question), unsafe_allow_html=True)
    
    # Answer in clean container
    st.markdown("""
    <div class="result-container">
        <h3>🧠 Expert Guidance</h3>
    </div>
    """, unsafe_allow_html=True)
    
    st.write(result["answer"])
    
    # Clean action buttons with UNIQUE KEYS
    st.markdown('<div class="clean-separator"></div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("🔄 Ask New Question", type="primary", use_container_width=True, key="result_new_question_button"):
            # Clear result to return to input mode
            st.session_state.ui_state.update({
                'current_result': None,
                'last_asked_question': "",
                'show_sources': False,
                'show_analytics': False,
                'question_counter': st.session_state.ui_state.get('question_counter', 0) + 1
            })
            st.rerun()
    
    with col2:
        if st.button("📄 Download Report", use_container_width=True, key="result_download_report_button"):
            generate_simple_report(result, question)
    
    with col3:
        if st.button("📚 Show Sources", use_container_width=True, key="result_show_sources_button"):
            st.session_state.ui_state['show_sources'] = not st.session_state.ui_state['show_sources']
    
    with col4:
        if st.button("📊 Analytics", use_container_width=True, key="result_show_analytics_button"):
            st.session_state.ui_state['show_analytics'] = not st.session_state.ui_state['show_analytics']
    
    # Progressive disclosure sections
    if st.session_state.ui_state['show_sources']:
        show_fixed_sources(result.get("sources", []))
    
    if st.session_state.ui_state['show_analytics']:
        show_simple_analytics(result)

def show_fixed_sources(sources):
    """FIXED source display with proper deduplication and unique keys"""
    st.markdown("---")
    st.markdown("### 📚 Reference Sources")
    
    if not sources:
        st.info("No specific source documents were referenced for this response.")
        return
    
    # Deduplicate sources using our helper function
    unique_sources = deduplicate_sources(sources)
    
    # Display unique sources
    for i, (doc_key, doc_info) in enumerate(unique_sources.items(), 1):
        title = clean_source_title(doc_info['title'], doc_info['source_file'], i)
        source_file = doc_info['source_file']
        chunks_found = doc_info['chunks_found']
        
        # Display the source
        st.write(f"**{i}.** {title}")
        
        if source_file:
            filename = os.path.basename(source_file)
            if chunks_found > 1:
                st.caption(f"📁 {filename} ({chunks_found} sections referenced)")
            else:
                st.caption(f"📁 {filename}")
        elif chunks_found > 1:
            st.caption(f"📄 {chunks_found} sections referenced")

def show_simple_analytics(result):
    """Simple analytics display with source information"""
    st.markdown("---")
    st.markdown("### 📊 Response Analytics")
    
    col1, col2 = st.columns(2)
    
    with col1:
        confidence = result.get("confidence_score", 0.0)
        if confidence >= 0.8:
            st.success(f"🟢 High Confidence ({confidence:.2f})")
        elif confidence >= 0.6:
            st.warning(f"🟡 Medium Confidence ({confidence:.2f})")
        else:
            st.error(f"🔴 Low Confidence ({confidence:.2f})")
        
        sources_count = len(result.get("sources", []))
        st.metric("Sources Consulted", sources_count)
        
        # Show unique sources count
        if sources_count > 0:
            unique_sources = deduplicate_sources(result.get("sources", []))
            unique_count = len(unique_sources)
            if unique_count != sources_count:
                st.caption(f"📄 {unique_count} unique documents")
    
    with col2:
        performance = result.get("performance", {})
        total_time = performance.get("total_response_time", 0)
        if total_time:
            st.metric("Response Time", f"{total_time:.1f}s")
        
        embedding_provider = result.get("routing_info", {}).get("embedding_provider", "unknown")
        if embedding_provider != "unknown":
            st.write(f"**Embedding Provider:** {embedding_provider.title()}")
        
        # Show retrieval time if available
        retrieval_time = performance.get("retrieval_time", 0)
        if retrieval_time:
            st.caption(f"⚡ Retrieval: {retrieval_time:.2f}s")

# ===== REPORT GENERATION =====
def generate_intelligent_filename(question, result):
    """Generate intelligent, descriptive filenames based on content"""
    
    # Get timestamp
    timestamp = time.strftime('%Y%m%d_%H%M%S')
    
    # Detect report type from question and result
    report_type = detect_report_type(question, result)
    
    # Generate descriptive name based on type
    if report_type == "ofsted_analysis":
        # Extract provider name if available
        provider_name = extract_provider_name(result.get("answer", ""))
        if provider_name:
            return f"Ofsted_Analysis_{provider_name}_{timestamp}.md"
        else:
            return f"Ofsted_Inspection_Analysis_{timestamp}.md"
    
    elif report_type == "policy_analysis":
        # Extract policy type if available
        policy_type = extract_policy_type(question)
        if policy_type:
            return f"Policy_Analysis_{policy_type}_{timestamp}.md"
        else:
            return f"Policy_Review_{timestamp}.md"
    
    elif report_type == "safeguarding":
        return f"Safeguarding_Assessment_{timestamp}.md"
    
    elif report_type == "compliance":
        return f"Regulatory_Compliance_Review_{timestamp}.md"
    
    elif report_type == "quality_assurance":
        return f"Quality_Assurance_Report_{timestamp}.md"
    
    elif report_type == "staff_development":
        return f"Staff_Development_Plan_{timestamp}.md"
    
    elif report_type == "inspection_prep":
        return f"Inspection_Preparation_Guide_{timestamp}.md"
    
    else:
        # Generate from question keywords
        clean_question = clean_question_for_filename(question)
        if clean_question:
            return f"Analysis_{clean_question}_{timestamp}.md"
        else:
            return f"Professional_Analysis_{timestamp}.md"

def detect_report_type(question, result):
    """Detect the type of report based on question and result content"""
    question_lower = question.lower()
    answer_lower = result.get("answer", "").lower()
    
    # Check for Ofsted analysis
    if any(keyword in question_lower for keyword in ["ofsted", "inspection report", "provider overview"]):
        return "ofsted_analysis"
    
    if any(keyword in answer_lower for keyword in ["provider name:", "inspection date:", "ratings by area"]):
        return "ofsted_analysis"
    
    # Check for policy analysis
    if any(keyword in question_lower for keyword in ["policy", "procedure", "compliance check"]):
        return "policy_analysis"
    
    # Check for safeguarding
    if any(keyword in question_lower for keyword in ["safeguarding", "child protection", "signs of safety"]):
        return "safeguarding"
    
    # Check for compliance
    if any(keyword in question_lower for keyword in ["regulatory", "compliance", "legal requirements"]):
        return "compliance"
    
    # Check for quality assurance
    if any(keyword in question_lower for keyword in ["quality", "monitoring", "improvement"]):
        return "quality_assurance"
    
    # Check for staff development
    if any(keyword in question_lower for keyword in ["staff training", "development", "supervision"]):
        return "staff_development"
    
    # Check for inspection preparation
    if any(keyword in question_lower for keyword in ["prepare", "preparation", "ready for"]):
        return "inspection_prep"
    
    return "general_analysis"

def extract_provider_name(answer_text):
    """Extract provider name from Ofsted analysis"""
    import re
    
    # Look for provider name patterns
    patterns = [
        r'\*\*Provider Name:\*\*\s*([^\n]+)',
        r'Provider Name:\s*([^\n]+)',
        r'\*\*([^*]+(?:Ltd|Limited|Care|Homes|Services))\*\*',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, answer_text, re.IGNORECASE)
        if match:
            provider = match.group(1).strip()
            # Clean the provider name for filename
            provider = re.sub(r'[^\w\s-]', '', provider)
            provider = re.sub(r'\s+', '_', provider)
            provider = provider[:30]  # Limit length
            return provider
    
    return None

def extract_policy_type(question):
    """Extract policy type from question"""
    import re
    
    # Common policy types
    policy_types = {
        'safeguarding': 'Safeguarding',
        'medication': 'Medication',
        'behaviour': 'Behaviour_Management',
        'education': 'Education',
        'health': 'Health_Safety',
        'admission': 'Admissions',
        'contact': 'Contact_Visits',
        'complaints': 'Complaints',
        'missing': 'Missing_Children',
        'restraint': 'Physical_Intervention'
    }
    
    question_lower = question.lower()
    for keyword, policy_type in policy_types.items():
        if keyword in question_lower:
            return policy_type
    
    return None

def clean_question_for_filename(question):
    """Clean question text to create readable filename"""
    import re
    
    # Take key words from question
    question = question.lower()
    
    # Remove common question words
    stop_words = ['what', 'how', 'when', 'where', 'why', 'who', 'can', 'should', 'would', 'could', 'is', 'are', 'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'from', 'as', 'into', 'onto', 'upon']
    
    # Extract meaningful words
    words = re.findall(r'\b\w+\b', question)
    meaningful_words = [word for word in words if word not in stop_words and len(word) > 2]
    
    # Take first 3-4 meaningful words
    key_words = meaningful_words[:4]
    
    if key_words:
        return '_'.join(key_words).title()
    
    return None

# Enhanced version of the generate_simple_report function with smart naming
def generate_simple_report(result, question):
    """Cleaner, more polished report generation"""
    
    # Helper function definitions (simplified and cleaner)
    def get_confidence_badge(confidence):
        """Create a simple confidence indicator"""
        if confidence >= 0.8:
            return f"High ({confidence:.0%})"
        elif confidence >= 0.6:
            return f"Medium ({confidence:.0%})"
        else:
            return f"Low ({confidence:.0%})"
    
    def create_executive_summary(answer_text):
        """Create a concise executive summary"""
        # Find first paragraph or first 2 sentences
        first_paragraph = answer_text.split('\n\n')[0] if answer_text else ""
        
        # If it's an Ofsted report, extract key info
        if "PROVIDER OVERVIEW" in answer_text and "Provider Name:" in answer_text:
            # Extract provider name and key finding
            lines = answer_text.split('\n')
            provider_line = next((line for line in lines if "Provider Name:" in line), "")
            provider = provider_line.replace("**Provider Name:**", "").strip() if provider_line else "Children's Home"
            
            # Find ratings
            requires_improvement = "Requires Improvement" in answer_text
            if requires_improvement:
                return f"Ofsted inspection of {provider} identifies areas requiring improvement in care provision and management oversight. Key actions needed include enhanced care planning, improved educational support, and strengthened management practices."
            else:
                return f"Ofsted inspection of {provider} shows generally positive outcomes with some areas for continued development."
        else:
            # General summary for non-Ofsted reports
            sentences = first_paragraph.split('. ')
            if len(sentences) >= 2:
                summary = '. '.join(sentences[:2]) + '.'
            else:
                summary = first_paragraph[:200] + "..." if len(first_paragraph) > 200 else first_paragraph
            
            return summary
    
    def clean_analysis_content(answer_text):
        """Clean the analysis content without over-formatting"""
        # Remove excessive emoji prefixes but keep some structure
        cleaned = answer_text.replace('## 🔍 ', '## ')
        cleaned = cleaned.replace('### 📌 ', '### ')
        cleaned = cleaned.replace('# 🔍 ', '## ')
        cleaned = cleaned.replace('# ## ', '## ')
        
        # Clean up any duplicate symbols
        cleaned = cleaned.replace('##\n## ', '## ')
        cleaned = cleaned.replace('\n\n\n', '\n\n')
        
        return cleaned
    
    def format_sources_simple(sources):
        """Simplified source formatting"""
        if not sources:
            return "No specific source documents were referenced for this analysis."
        
        # Deduplicate sources
        unique_sources = deduplicate_sources(sources)
        
        sources_text = "**Sources Referenced:**\n\n"
        
        for i, (doc_key, doc_info) in enumerate(unique_sources.items(), 1):
            title = clean_source_title(doc_info['title'], doc_info['source_file'], i)
            word_count = doc_info['word_count']
            chunks = doc_info['chunks_found']
            
            sources_text += f"{i}. **{title}**\n"
            if word_count > 0:
                content_desc = f"   - {word_count:,} words"
                if chunks > 1:
                    content_desc += f" ({chunks} sections referenced)"
                sources_text += content_desc + "\n"
            sources_text += "\n"
        
        return sources_text
    
    def detect_report_type(question, result):
        """Detect the type of report"""
        question_lower = question.lower()
        answer_lower = result.get("answer", "").lower()
        
        if any(keyword in question_lower for keyword in ["ofsted", "inspection report"]):
            return "ofsted_analysis"
        if any(keyword in answer_lower for keyword in ["provider name:", "inspection date:", "ratings by area"]):
            return "ofsted_analysis"
        if any(keyword in question_lower for keyword in ["policy", "procedure"]):
            return "policy_analysis"
        if any(keyword in question_lower for keyword in ["safeguarding", "child protection"]):
            return "safeguarding"
        
        return "general_analysis"
    
    def generate_filename(question, result):
        """Generate clean, descriptive filename"""
        import re
        
        timestamp = time.strftime('%Y%m%d_%H%M%S')
        report_type = detect_report_type(question, result)
        
        if report_type == "ofsted_analysis":
            # Try to extract provider name
            answer = result.get("answer", "")
            provider_match = re.search(r'\*\*Provider Name:\*\*\s*([^\n]+)', answer)
            if provider_match:
                provider = provider_match.group(1).strip()
                provider = re.sub(r'[^\w\s-]', '', provider)
                provider = re.sub(r'\s+', '_', provider)[:25]
                return f"Ofsted_Analysis_{provider}_{timestamp}.md"
            return f"Ofsted_Analysis_{timestamp}.md"
        
        elif report_type == "policy_analysis":
            return f"Policy_Analysis_{timestamp}.md"
        elif report_type == "safeguarding":
            return f"Safeguarding_Report_{timestamp}.md"
        else:
            # Extract key words from question
            words = re.findall(r'\b\w{4,}\b', question.lower())
            key_words = [w for w in words if w not in ['what', 'how', 'when', 'where', 'analysis', 'report']][:3]
            if key_words:
                desc = '_'.join(key_words).title()
                return f"Analysis_{desc}_{timestamp}.md"
            return f"Professional_Analysis_{timestamp}.md"
    
    def get_report_title(question, result):
        """Generate clean report title"""
        report_type = detect_report_type(question, result)
        
        titles = {
            "ofsted_analysis": "Ofsted Inspection Analysis",
            "policy_analysis": "Policy & Procedures Review",
            "safeguarding": "Safeguarding Assessment",
            "general_analysis": "Professional Analysis Report"
        }
        
        return titles.get(report_type, "Analysis Report")
    
    # MAIN FUNCTION LOGIC
    sources = result.get("sources", [])
    confidence = result.get("confidence_score", 0.0)
    performance = result.get("performance", {})
    metadata = result.get("metadata", {})
    
    report_title = get_report_title(question, result)
    executive_summary = create_executive_summary(result["answer"])
    cleaned_analysis = clean_analysis_content(result["answer"])
    
    # Create clean, professional report
    report_content = f"""# Children's Home Management System
## {report_title}

---

**Generated:** {time.strftime('%d %B %Y at %H:%M')}  
**Performance Mode:** {st.session_state.get('performance_mode', 'balanced').title()}  
**Confidence Level:** {get_confidence_badge(confidence)}  
**Report ID:** RPT-{time.strftime('%Y%m%d-%H%M%S')}

---

## Executive Summary

{executive_summary}

---

## Question

> {question}

---

## Analysis

{cleaned_analysis}

---

## Sources

{format_sources_simple(sources)}

---

## Technical Information

**AI Model:** {metadata.get('llm_used', 'Unknown')}  
**Response Time:** {performance.get('total_response_time', 0):.1f} seconds  
**Sources Consulted:** {len(sources)}  
**Session:** {st.session_state.get('tester_id', 'anonymous')}

---

*This report was generated by the Children's Home Management System on {time.strftime('%A, %d %B %Y')}.*

**Important:** This AI-generated analysis should be reviewed by qualified professionals for final decision-making.
"""
    
    # Generate filename and download button
    filename = generate_filename(question, result)
    
    st.download_button(
        label="📄 Download Report",
        data=report_content,
        file_name=filename,
        mime="text/markdown",
        use_container_width=True,
        key="download_clean_report_button",
        help="Download professional analysis report"
    )

def generate_report_title(question, result):
    """Generate intelligent report title"""
    
    report_type = detect_report_type(question, result)
    
    title_map = {
        "ofsted_analysis": "Ofsted Inspection Analysis Report",
        "policy_analysis": "Policy & Procedures Analysis Report",
        "safeguarding": "Safeguarding Assessment Report",
        "compliance": "Regulatory Compliance Review",
        "quality_assurance": "Quality Assurance Report",
        "staff_development": "Staff Development Analysis",
        "inspection_prep": "Inspection Preparation Guide",
        "general_analysis": "Professional Analysis Report"
    }
    
    return title_map.get(report_type, "Professional Analysis Report")

# ===== SYSTEM HEALTH DISPLAY =====
def display_system_health():
    """Display system health and performance information"""
    st.subheader("⚙️ System Health & Performance")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 🤖 AI Models Status")
        
        if st.session_state.get('rag_system'):
            try:
                health = st.session_state.rag_system.get_system_health()
                models = health.get('llm_models', {})
                
                st.write(f"✅ Available Models: {len(models.get('available_models', []))}")
                for model in models.get('available_models', []):
                    st.write(f"  • {model}")
                    
            except Exception as e:
                st.error(f"❌ Health check failed: {e}")
        else:
            st.error("❌ RAG System not available")
        
        st.markdown("### 🔑 API Configuration")
        openai_key = bool(os.environ.get("OPENAI_API_KEY"))
        google_key = bool(os.environ.get("GOOGLE_API_KEY"))
        
        st.write(f"{'✅' if openai_key else '❌'} OpenAI API")
        st.write(f"{'✅' if google_key else '❌'} Google API")
    
    with col2:
        st.markdown("### 📊 Performance Metrics")
        
        if hasattr(st.session_state.rag_system, 'get_performance_stats'):
            try:
                perf_stats = st.session_state.rag_system.get_performance_stats()
                if isinstance(perf_stats, dict) and perf_stats.get('total_queries', 0) > 0:
                    st.metric("Total Queries", perf_stats['total_queries'])
                    st.metric("Avg Response Time", f"{perf_stats.get('avg_response_time', 0):.1f}s")
                    st.metric("Success Rate", f"{(perf_stats.get('successful_queries', 0)/perf_stats.get('total_queries', 1)*100):.1f}%")
                else:
                    st.info("No performance data available yet")
            except Exception as e:
                st.warning(f"Performance stats unavailable: {e}")
        
        if SYSTEM_MONITOR:
            st.markdown("### 💾 System Resources")
            try:
                memory_percent = psutil.virtual_memory().percent
                st.metric("Memory Usage", f"{memory_percent:.1f}%")
            except Exception as e:
                st.warning(f"Memory monitoring unavailable: {e}")
        
        st.markdown("### ⏱️ Session Info")
        session_start = st.session_state.get('session_start', time.time())
        session_duration = int((time.time() - session_start) / 60)
        st.metric("Session Duration", f"{session_duration} min")
    
    with st.expander("🔍 Detailed System Information"):
        if st.session_state.get('rag_system'):
            try:
                health = st.session_state.rag_system.get_system_health()
                st.json(health)
            except Exception as e:
                st.error(f"Detailed health check failed: {e}")

# ===== SESSION MANAGEMENT =====
def show_session_summary():
    """Enhanced session summary with performance stats"""
    with st.sidebar:
        if st.session_state.get('activity_log'):
            st.markdown("---")
            st.markdown("### 📊 Session Summary")
            
            activities = st.session_state.get('activity_log', [])
            question_count = len([a for a in activities if a['action'] in ['unified_request_submitted', 'question_submitted']])
            feedback_count = len([a for a in activities if a['action'] == 'feedback_submitted'])
            
            st.write(f"**Questions asked:** {question_count}")
            st.write(f"**Feedback items:** {feedback_count}")
            
            session_minutes = int((time.time() - st.session_state.get('session_start', time.time())) / 60)
            st.write(f"**Session time:** {session_minutes} minutes")
            
            # Performance stats if available
            if hasattr(st.session_state.rag_system, 'get_performance_stats'):
                try:
                    perf_stats = st.session_state.rag_system.get_performance_stats()
                    if isinstance(perf_stats, dict) and perf_stats.get('total_queries', 0) > 0:
                        st.markdown("**⚡ Performance:**")
                        avg_time = perf_stats.get('avg_response_time', 0)
                        success_rate = perf_stats.get('successful_queries', 0) / perf_stats.get('total_queries', 1) * 100
                        st.write(f"Avg response: {avg_time:.1f}s")
                        st.write(f"Success rate: {success_rate:.1f}%")
                except Exception:
                    pass
            
            with st.expander("📋 Activity Log"):
                for activity in activities[-10:]:
                    st.text(f"{activity['timestamp']}: {activity['action']}")

# ===== MAIN APPLICATION =====
def main():
    """Main application function with unified interface"""
    setup_sqlite_with_clean_logging()

    if not require_authentication():
        st.stop()
    
    if not setup_environment_variables():
        st.error("❌ Failed to configure API keys")
        st.info("💡 Check your .streamlit/secrets.toml file")
        with st.expander("🔧 Configuration Help"):
            st.code("""
# .streamlit/secrets.toml
[api_keys]
openai = "your-openai-api-key"
google = "your-google-api-key"
            """)
        st.stop()
    
    if 'rag_system' not in st.session_state:
        with st.spinner("Initializing system..."):
            st.session_state.rag_system = initialize_rag_system()
    
    if st.session_state.rag_system is None:
        st.error("❌ System initialization failed")
        st.info("💡 Please refresh the page or contact support")
        st.stop()
    
    # CRITICAL: Ensure UI state is initialized early
    ensure_ui_state()
    
    # Load CSS once
    st.markdown(get_app_css(), unsafe_allow_html=True)
    
    # Beta testing banner
    st.markdown("""
    <div class="testing-banner">
    🧪 BETA TESTING VERSION - Your feedback helps improve this system!
    </div>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown('<div class="main-header">', unsafe_allow_html=True)
    st.title("🏠 Children's Home Management System")
    st.markdown("*Ask questions, upload files, or analyze images - all in one place*")
    
    tester_info = st.session_state.get('tester_info', {})
    if tester_info:
        st.markdown(f"*Welcome, {tester_info.get('name', 'Beta Tester')}!*")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    add_performance_mode_selector()
    
    # Feedback mechanism in sidebar
    with st.sidebar:
        st.markdown("### 📝 Beta Feedback")
        
        feedback_type = st.selectbox(
            "Feedback Type:",
            ["💡 Suggestion", "🐛 Bug Report", "👍 Positive", "❓ Question"],
            key="feedback_type_select_widget"
        )
        
        feedback_text = st.text_area(
            "Your Feedback:",
            placeholder="Share thoughts, report bugs, or suggest improvements...",
            height=100,
            key="feedback_text_input_widget"
        )
        
        if st.button("📤 Submit Feedback", type="secondary", key="feedback_submit_button"):
            if feedback_text.strip():
                log_tester_activity(
                    st.session_state.get('tester_id', 'unknown'),
                    "feedback_submitted",
                    f"{feedback_type}: {feedback_text[:100]}..."
                )
                st.success("✅ Feedback submitted!")
            else:
                st.warning("Please enter feedback before submitting")
    
    # Debug controls
    with st.sidebar:
        if st.session_state.get('tester_id') in ['DEMO001', 'TEST001']:
            st.markdown("---")
            st.markdown("### 🔧 Debug Controls")
            
            current_debug = os.getenv('DEBUG_MODE', 'false').lower() == 'true'
            debug_mode = st.checkbox(
                "🔍 Debug Mode", 
                value=current_debug,
                help="Show detailed system logs in terminal",
                key="debug_mode_checkbox_widget"
            )
            
            if debug_mode != current_debug:
                if debug_mode:
                    os.environ['DEBUG_MODE'] = 'true'
                    st.success("🔍 Debug logging enabled")
                else:
                    os.environ['DEBUG_MODE'] = 'false'
                    st.info("🔇 Debug logging disabled")
            
            if debug_mode:
                st.caption("📝 Debug messages will appear in terminal")
            else:
                st.caption("🔇 Only critical events logged")
    
    show_session_summary()
    
    # MAIN CONTENT LOGIC: Show EITHER input interface OR results (not both)
    if st.session_state.ui_state.get('current_result'):
        # RESULTS MODE: Show only the clean result display
        show_clean_result_display()
    else:
        # INPUT MODE: Show only the input interface
        show_unified_input_interface()
    
    # System health in expandable section (always show)
    with st.expander("⚙️ System Health & Performance", expanded=False):
        display_system_health()

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; font-size: 0.9em; padding: 1rem;">
    🧪 <strong>Thank you for beta testing!</strong> Your feedback helps improve this system.<br>
    Use the feedback form in the sidebar to report issues or suggestions.
    </div>
    """, unsafe_allow_html=True)

# ===== MAIN EXECUTION =====
if __name__ == "__main__":
    if 'page_loads' not in st.session_state:
        st.session_state['page_loads'] = 0
        log_tester_activity(
            st.session_state.get('tester_id', 'unknown'),
            "session_started",
            "Application loaded"
        )
    
    st.session_state['page_loads'] += 1
    main()
